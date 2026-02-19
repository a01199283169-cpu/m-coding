import sys
import os
import re
from pathlib import Path
from datetime import datetime

import pdfplumber
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QTextEdit, QFileDialog, QMessageBox,
    QProgressBar, QFrame,
)
from PyQt6.QtCore import QThread, pyqtSignal, Qt
from PyQt6.QtGui import QFont


# ──────────────────────────────────────────────────────────
# Worker: PDF 분석
# ──────────────────────────────────────────────────────────
class PDFAnalysisWorker(QThread):
    log_signal      = pyqtSignal(str)
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(list)
    error_signal    = pyqtSignal(str)

    def __init__(self, folder_path):
        super().__init__()
        self.folder_path = folder_path

    # ── 메인 실행 ──────────────────────────────────────────
    def run(self):
        try:
            pdf_files = sorted(Path(self.folder_path).glob("*.pdf"))
            if not pdf_files:
                self.error_signal.emit("선택한 폴더에 PDF 파일이 없습니다.")
                return

            self.log_signal.emit(f"총 {len(pdf_files)}개의 PDF 파일을 발견했습니다.\n")
            all_data = []

            for i, pdf_path in enumerate(pdf_files):
                self.log_signal.emit(f"[{i+1}/{len(pdf_files)}] {pdf_path.name} 처리 중...")
                try:
                    rows = self.extract_from_pdf(pdf_path)
                    all_data.extend(rows)
                    self.log_signal.emit(f"  → {len(rows)}건 안건 추출 완료")
                except Exception as e:
                    self.log_signal.emit(f"  → 오류: {e}")

                self.progress_signal.emit(int((i + 1) / len(pdf_files) * 80))

            self.log_signal.emit(f"\n총 {len(all_data)}건 데이터 추출 완료.")
            self.finished_signal.emit(all_data)

        except Exception as e:
            self.error_signal.emit(f"분석 중 오류: {e}")

    # ── PDF 1개 처리 ───────────────────────────────────────
    def extract_from_pdf(self, pdf_path: Path) -> list:
        results = []
        with pdfplumber.open(pdf_path) as pdf:
            full_text = ""
            all_tables = []
            for page in pdf.pages:
                full_text += (page.extract_text() or "") + "\n"
                tables = page.extract_tables()
                if tables:
                    all_tables.extend(tables)

        mtg_name = self._meeting_name(full_text)
        date_str = self._meeting_date(full_text)

        # 테이블 우선 파싱
        agendas = self._parse_tables(all_tables)
        # 테이블 실패 시 텍스트 파싱
        if not agendas:
            agendas = self._parse_text(full_text)

        if agendas:
            for ag in agendas:
                results.append({
                    "파일명":        pdf_path.name,
                    "회의명":        mtg_name,
                    "회의일":        date_str,
                    "안건명":        ag.get("name", ""),
                    "안건 주요사항": ag.get("content", ""),
                    "의결결과":      ag.get("result", ""),
                    "비고":          ag.get("remark", ""),
                })
        else:
            results.append({
                "파일명":        pdf_path.name,
                "회의명":        mtg_name,
                "회의일":        date_str,
                "안건명":        "(안건 추출 불가)",
                "안건 주요사항": "",
                "의결결과":      "",
                "비고":          "",
            })
        return results

    # ── 회의명 추출 ────────────────────────────────────────
    def _meeting_name(self, text: str) -> str:
        patterns = [
            r'제\s*\d+\s*차\s*(?:정기|임시)?\s*입주자대표\s*회의',
            r'(?:정기|임시)\s*입주자대표\s*회의',
            r'제\s*\d+\s*차\s*(?:정기|임시)?\s*이사회',
            r'입주자대표\s*회의',
            r'이사회',
        ]
        for p in patterns:
            m = re.search(p, text)
            if m:
                return re.sub(r'\s+', ' ', m.group(0)).strip()
        return ""

    # ── 회의일 추출 ────────────────────────────────────────
    def _meeting_date(self, text: str) -> str:
        # 키워드 주변 날짜 우선
        ctx = re.search(
            r'(?:일\s*시|회의일시|개최일|날\s*짜)[^\n]*?'
            r'(\d{4})\s*[년./\-]\s*(\d{1,2})\s*[월./\-]\s*(\d{1,2})',
            text,
        )
        if ctx:
            return f"{ctx.group(1)}년 {ctx.group(2)}월 {ctx.group(3)}일"

        m = re.search(
            r'(\d{4})\s*[년./\-]\s*(\d{1,2})\s*[월./\-]\s*(\d{1,2})', text
        )
        if m:
            return f"{m.group(1)}년 {m.group(2)}월 {m.group(3)}일"
        return ""

    # ── 테이블 파싱 ────────────────────────────────────────
    AGENDA_KW  = {"안건", "안건명", "심의안건", "의결안건", "제목"}
    RESULT_KW  = {"의결결과", "결과", "처리결과", "의결사항", "처리사항"}
    CONTENT_KW = {"주요내용", "내용", "주요사항", "심의내용", "검토내용", "처리내용"}
    REMARK_KW  = {"비고", "특이사항", "기타사항"}
    RESOLVE_KW = ["원안가결", "수정가결", "부결", "보류", "가결", "승인", "의결"]

    def _parse_tables(self, tables: list) -> list:
        agendas = []
        for table in tables:
            if not table:
                continue
            col_ag, col_rs, col_ct, col_rk = -1, -1, -1, -1
            hdr_idx = -1

            # 헤더 행 탐색 (처음 5행)
            for ri, row in enumerate(table[:5]):
                if not row:
                    continue
                joined = " ".join(str(c) for c in row if c)
                if any(k in joined for k in self.AGENDA_KW | self.RESULT_KW):
                    hdr_idx = ri
                    for ci, cell in enumerate(row):
                        s = str(cell) if cell else ""
                        if any(k in s for k in self.AGENDA_KW):
                            col_ag = ci
                        if any(k in s for k in self.RESULT_KW):
                            col_rs = ci
                        if any(k in s for k in self.CONTENT_KW):
                            col_ct = ci
                        if any(k in s for k in self.REMARK_KW):
                            col_rk = ci
                    break

            if hdr_idx < 0 or col_ag < 0:
                continue

            for row in table[hdr_idx + 1:]:
                if not row:
                    continue
                def _get(ci):
                    return str(row[ci]).strip() if 0 <= ci < len(row) and row[ci] else ""

                name    = _get(col_ag)
                result  = _get(col_rs)
                content = _get(col_ct)
                remark  = _get(col_rk)

                if not name or name in {"None", "-"}:
                    continue
                if not result:
                    result = self._resolve_in_row(row)

                agendas.append({"name": name, "content": content,
                                "result": result, "remark": remark})

        return agendas

    def _resolve_in_row(self, row) -> str:
        text = " ".join(str(c) for c in row if c)
        for kw in self.RESOLVE_KW:
            if kw in text:
                m = re.search(kw + r'[^\s\n]*', text)
                return m.group(0) if m else kw
        return ""

    # ── 텍스트 파싱 ────────────────────────────────────────
    def _parse_text(self, text: str) -> list:
        agendas = []
        lines = text.split("\n")

        agenda_re = [
            re.compile(r'(?:제|안건)?\s*(\d+)\s*(?:호|번)?\s*안건\s*[:\.\s]\s*(.{3,})'),
            re.compile(r'안건\s*(\d+)\s*[:\.\s]\s*(.{3,})'),
            re.compile(r'^(\d+)\.\s+(.{5,60})$'),
        ]

        for i, line in enumerate(lines):
            line = line.strip()
            for pat in agenda_re:
                m = pat.match(line)
                if m:
                    name   = m.group(2).strip()
                    result = ""
                    content_lines = []

                    for j in range(i + 1, min(i + 12, len(lines))):
                        sl = lines[j].strip()
                        if not sl:
                            continue
                        # 의결결과 키워드 발견 시 캡처 후 중단
                        hit = False
                        for kw in self.RESOLVE_KW:
                            if kw in sl:
                                rm = re.search(kw + r'[^\n]*', sl)
                                result = rm.group(0).strip() if rm else kw
                                hit = True
                                break
                        if hit:
                            break
                        # 다음 안건 시작 시 중단
                        if any(p.match(sl) for p in agenda_re):
                            break
                        content_lines.append(sl)

                    content = ' '.join(content_lines[:3])  # 최대 3줄
                    agendas.append({"name": name, "content": content,
                                    "result": result, "remark": ""})
                    break

        return agendas


# ──────────────────────────────────────────────────────────
# Worker: Excel / Word 저장
# ──────────────────────────────────────────────────────────
class SaveWorker(QThread):
    log_signal      = pyqtSignal(str)
    finished_signal = pyqtSignal(str, str)
    error_signal    = pyqtSignal(str)

    def __init__(self, data: list, save_folder: str):
        super().__init__()
        self.data        = data
        self.save_folder = save_folder

    def run(self):
        try:
            excel_path = self._save_excel()
            word_path  = self._save_word()
            self.finished_signal.emit(excel_path, word_path)
        except Exception as e:
            self.error_signal.emit(f"저장 중 오류: {e}")

    # ── Excel ──────────────────────────────────────────────
    def _save_excel(self) -> str:
        self.log_signal.emit("\n[Excel] 저장 중...")
        df = pd.DataFrame(self.data)
        path = str(Path(self.save_folder) / "회의록_분석결과.xlsx")

        COLS   = ["파일명", "회의명", "회의일", "안건명", "안건 주요사항", "의결결과", "비고"]
        WIDTHS = [26, 22, 14, 32, 42, 14, 12]

        with pd.ExcelWriter(path, engine="xlsxwriter") as writer:
            df.to_excel(writer, sheet_name="회의록", index=False)
            wb  = writer.book
            ws  = writer.sheets["회의록"]

            # ── 서식 정의 ──
            hdr_fmt = wb.add_format({
                "bold": True, "font_size": 11,
                "bg_color": "#2E75B6", "font_color": "#FFFFFF",
                "align": "center", "valign": "vcenter", "border": 1,
            })
            row_odd = wb.add_format({
                "font_size": 10, "valign": "vcenter",
                "border": 1, "text_wrap": True,
            })
            row_even = wb.add_format({
                "font_size": 10, "valign": "vcenter",
                "border": 1, "text_wrap": True, "bg_color": "#EBF3FB",
            })
            pass_fmt = wb.add_format({
                "font_size": 10, "align": "center", "valign": "vcenter",
                "border": 1, "bold": True,
                "font_color": "#196B24", "bg_color": "#E2EFDA",
            })
            fail_fmt = wb.add_format({
                "font_size": 10, "align": "center", "valign": "vcenter",
                "border": 1, "bold": True,
                "font_color": "#9C0006", "bg_color": "#FFC7CE",
            })

            # 헤더
            for ci, (col, w) in enumerate(zip(COLS, WIDTHS)):
                ws.write(0, ci, col, hdr_fmt)
                ws.set_column(ci, ci, w)

            # 데이터
            res_ci = COLS.index("의결결과")
            for ri in range(len(df)):
                base = row_even if ri % 2 == 0 else row_odd
                for ci, col in enumerate(COLS):
                    val = df.iloc[ri][col] if col in df.columns else ""
                    if ci == res_ci:
                        vs = str(val)
                        if any(k in vs for k in ["가결", "승인"]):
                            ws.write(ri + 1, ci, val, pass_fmt)
                        elif any(k in vs for k in ["부결", "보류"]):
                            ws.write(ri + 1, ci, val, fail_fmt)
                        else:
                            ws.write(ri + 1, ci, val, base)
                    else:
                        ws.write(ri + 1, ci, val, base)
                ws.set_row(ri + 1, 20)

            ws.set_row(0, 25)
            ws.freeze_panes(1, 0)

        self.log_signal.emit(f"[Excel] 저장 완료: {path}")
        return path

    # ── Word ───────────────────────────────────────────────
    def _save_word(self) -> str:
        self.log_signal.emit("\n[Word] 저장 중...")
        doc  = Document()
        path = str(Path(self.save_folder) / "회의록_분석결과.docx")

        # 페이지 여백
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(3.0)
            sec.right_margin  = Cm(3.0)

        # 제목
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title.add_run("입주자대표회의 회의록 분석 보고서")
        tr.bold           = True
        tr.font.size      = Pt(18)
        tr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        # 작성일
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        sr.font.size      = Pt(10)
        sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        # 요약
        df = pd.DataFrame(self.data)
        smry = doc.add_paragraph()
        smr  = smry.add_run(
            f"■ 분석 요약: 총 {len(df)}건의 안건  /  {df['파일명'].nunique()}개 파일"
        )
        smr.bold      = True
        smr.font.size = Pt(11)

        doc.add_paragraph()

        # 표
        HEADERS   = ["파일명", "회의명", "회의일", "안건명", "안건 주요사항", "의결결과", "비고"]
        COL_W_CM  = [3.5, 3.0, 2.8, 3.8, 5.0, 2.4, 1.5]

        table = doc.add_table(rows=1, cols=len(HEADERS))
        table.style = "Table Grid"

        # 헤더 행
        for ci, h in enumerate(HEADERS):
            cell = table.rows[0].cells[ci]
            cell.text = h
            self._cell_shading(cell, "2E75B6")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold            = True
                    run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size       = Pt(10)

        # 데이터 행
        res_ci = HEADERS.index("의결결과")
        for ri, row_data in enumerate(self.data):
            row_cells = table.add_row().cells
            values = [row_data.get(h, "") for h in HEADERS]

            bg = "EBF3FB" if ri % 2 == 0 else None

            for ci, val in enumerate(values):
                cell = row_cells[ci]
                cell.text = str(val) if val else ""
                if bg:
                    self._cell_shading(cell, bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if ci == res_ci:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run.bold = True
                            vs = str(val)
                            if any(k in vs for k in ["가결", "승인"]):
                                run.font.color.rgb = RGBColor(0x19, 0x6B, 0x24)
                            elif any(k in vs for k in ["부결", "보류"]):
                                run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

        # 열 너비
        for ri_t, row in enumerate(table.rows):
            for ci, w in enumerate(COL_W_CM):
                row.cells[ci].width = Cm(w)

        doc.save(path)
        self.log_signal.emit(f"[Word] 저장 완료: {path}")
        return path

    @staticmethod
    def _cell_shading(cell, fill_hex: str):
        """표 셀 배경색 적용"""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        tcPr.append(shd)


# ──────────────────────────────────────────────────────────
# 메인 윈도우
# ──────────────────────────────────────────────────────────
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.folder_path    = None
        self.extracted_data = []
        self._build_ui()

    def _build_ui(self):
        self.setWindowTitle("입주자대표회의 회의록 분석기")
        self.setMinimumSize(760, 620)
        self.resize(860, 720)

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        # 제목
        lbl_title = QLabel("입주자대표회의 회의록 분석기")
        f = QFont(); f.setPointSize(16); f.setBold(True)
        lbl_title.setFont(f)
        lbl_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl_title.setStyleSheet("color:#2E75B6; padding:8px;")
        layout.addWidget(lbl_title)

        # 구분선
        hr = QFrame(); hr.setFrameShape(QFrame.Shape.HLine)
        hr.setFrameShadow(QFrame.Shadow.Sunken)
        layout.addWidget(hr)

        # 폴더 선택
        row1 = QHBoxLayout()
        self.lbl_folder = QLabel("폴더를 선택해주세요")
        self.lbl_folder.setStyleSheet(
            "border:1px solid #ccc; padding:8px; border-radius:4px; background:#f9f9f9;"
        )
        self.lbl_folder.setMinimumHeight(36)
        row1.addWidget(self.lbl_folder)

        self.btn_select = QPushButton("회의자료 폴더 선택")
        self._style_btn(self.btn_select, "#2E75B6", "#1F5C94")
        self.btn_select.setMinimumWidth(165)
        self.btn_select.setMinimumHeight(36)
        self.btn_select.clicked.connect(self._select_folder)
        row1.addWidget(self.btn_select)
        layout.addLayout(row1)

        # 진행 바
        self.progress = QProgressBar()
        self.progress.setValue(0)
        self.progress.setStyleSheet(
            "QProgressBar{border:1px solid #ccc; border-radius:4px;"
            "text-align:center; height:22px;}"
            "QProgressBar::chunk{background:#2E75B6; border-radius:3px;}"
        )
        layout.addWidget(self.progress)

        # 로그 창
        lbl_log = QLabel("처리 로그:")
        lbl_log.setStyleSheet("font-weight:bold; color:#333;")
        layout.addWidget(lbl_log)

        self.log_box = QTextEdit()
        self.log_box.setReadOnly(True)
        self.log_box.setMinimumHeight(300)
        self.log_box.setStyleSheet(
            "QTextEdit{border:1px solid #ccc; border-radius:4px;"
            "background:#1E1E1E; color:#D4D4D4;"
            "font-family:Consolas,'Malgun Gothic',monospace;"
            "font-size:12px; padding:5px;}"
        )
        layout.addWidget(self.log_box)

        # 버튼 행
        row2 = QHBoxLayout(); row2.setSpacing(10)

        self.btn_run = QPushButton("분석 및 저장")
        self._style_btn(self.btn_run, "#107C41", "#0B5E31",
                        disabled_bg="#A0A0A0")
        self.btn_run.setMinimumHeight(44)
        self.btn_run.setStyleSheet(
            self.btn_run.styleSheet() +
            "QPushButton{font-size:14px; font-weight:bold;}"
        )
        self.btn_run.setEnabled(False)
        self.btn_run.clicked.connect(self._run_analysis)
        row2.addWidget(self.btn_run)

        btn_clr = QPushButton("로그 지우기")
        self._style_btn(btn_clr, "#666666", "#444444")
        btn_clr.setMinimumHeight(44)
        btn_clr.setMaximumWidth(120)
        btn_clr.clicked.connect(self.log_box.clear)
        row2.addWidget(btn_clr)

        layout.addLayout(row2)
        self.statusBar().showMessage("준비")

    # ── 스타일 헬퍼 ───────────────────────────────────────
    @staticmethod
    def _style_btn(btn, bg, hover, disabled_bg=None):
        disabled_css = (
            f"QPushButton:disabled{{background:{disabled_bg}; color:#E0E0E0;}}"
            if disabled_bg else ""
        )
        btn.setStyleSheet(
            f"QPushButton{{background:{bg}; color:white; border:none;"
            f"border-radius:4px; padding:0 14px;}}"
            f"QPushButton:hover{{background:{hover};}}"
            f"QPushButton:pressed{{background:{hover};}}"
            + disabled_css
        )

    # ── 슬롯 ──────────────────────────────────────────────
    def _select_folder(self):
        folder = QFileDialog.getExistingDirectory(
            self, "회의자료 폴더 선택", "",
            QFileDialog.Option.ShowDirsOnly,
        )
        if not folder:
            return
        self.folder_path = folder
        self.lbl_folder.setText(folder)
        self.btn_run.setEnabled(True)
        n = len(list(Path(folder).glob("*.pdf")))
        self._log(f"폴더 선택됨: {folder}")
        self._log(f"PDF 파일 {n}개 발견")

    def _run_analysis(self):
        if not self.folder_path:
            QMessageBox.warning(self, "경고", "먼저 폴더를 선택해주세요.")
            return

        self.btn_run.setEnabled(False)
        self.btn_select.setEnabled(False)
        self.progress.setValue(0)
        self._log("\n" + "=" * 50)
        self._log("PDF 분석을 시작합니다...")

        self._worker_pdf = PDFAnalysisWorker(self.folder_path)
        self._worker_pdf.log_signal.connect(self._log)
        self._worker_pdf.progress_signal.connect(self.progress.setValue)
        self._worker_pdf.finished_signal.connect(self._on_pdf_done)
        self._worker_pdf.error_signal.connect(self._on_error)
        self._worker_pdf.start()

    def _on_pdf_done(self, data: list):
        self.extracted_data = data
        if not data:
            self._log("\n추출된 데이터가 없습니다.")
            self._reset_buttons()
            return

        self._log(f"\n데이터 추출 완료. 파일 저장을 시작합니다...")

        self._worker_save = SaveWorker(data, self.folder_path)
        self._worker_save.log_signal.connect(self._log)
        self._worker_save.finished_signal.connect(self._on_save_done)
        self._worker_save.error_signal.connect(self._on_error)
        self._worker_save.start()

    def _on_save_done(self, excel_path: str, word_path: str):
        self.progress.setValue(100)
        self._log("\n" + "=" * 50)
        self._log("모든 작업이 완료되었습니다!")
        self._log(f"  Excel : {excel_path}")
        self._log(f"  Word  : {word_path}")
        self.statusBar().showMessage("완료")
        self._reset_buttons()
        QMessageBox.information(
            self, "완료",
            f"분석 및 저장이 완료되었습니다.\n\n"
            f"Excel: {excel_path}\n"
            f"Word : {word_path}",
        )

    def _on_error(self, msg: str):
        self._log(f"\n[오류] {msg}")
        self.statusBar().showMessage("오류 발생")
        self._reset_buttons()
        QMessageBox.critical(self, "오류", msg)

    def _log(self, text: str):
        self.log_box.append(text)
        self.log_box.verticalScrollBar().setValue(
            self.log_box.verticalScrollBar().maximum()
        )

    def _reset_buttons(self):
        self.btn_run.setEnabled(True)
        self.btn_select.setEnabled(True)


# ──────────────────────────────────────────────────────────
# 진입점
# ──────────────────────────────────────────────────────────
def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    win = MainWindow()
    win.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
