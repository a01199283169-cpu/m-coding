"""
회의록 분석기 - Flask 웹 서버
PyQt6 앱의 핵심 로직을 그대로 재활용, UI만 웹으로 교체
"""
import os
import re
import uuid
import json
import queue
import threading
from pathlib import Path
from datetime import datetime

from flask import Flask, render_template, request, jsonify, Response, send_file, stream_with_context
import pdfplumber
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)

# job_id → {queue, status, output_dir}
jobs: dict = {}


# ── PDF 분석 로직 ──────────────────────────────────────────────

AGENDA_KW  = {"안건", "안건명", "심의안건", "의결안건", "제목"}
RESULT_KW  = {"의결결과", "결과", "처리결과", "의결사항", "처리사항"}
CONTENT_KW = {"주요내용", "내용", "주요사항", "심의내용", "검토내용", "처리내용"}
REMARK_KW  = {"비고", "특이사항", "기타사항"}
RESOLVE_KW = ["원안가결", "수정가결", "부결", "보류", "가결", "승인", "의결"]


def meeting_name(text: str) -> str:
    """PDF 전문에서 회의명을 추출"""
    patterns = [
        r'제\s*\d+\s*차\s*(?:정기|임시)?\s*입주자대표\s*회의',
        r'(?:정기|임시)\s*입주자대표\s*회의',
        r'제\s*\d+\s*차\s*(?:정기|임시)?\s*이사회',
        r'입주자대표\s*회의',
        r'이사회',
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return re.sub(r'\s+', ' ', m.group(0)).strip()
    return ""


def meeting_date(text: str) -> str:
    ctx = re.search(
        r'(?:일\s*시|회의일시|개최일|날\s*짜)[^\n]*?'
        r'(\d{4})\s*[년./\-]\s*(\d{1,2})\s*[월./\-]\s*(\d{1,2})', text
    )
    if ctx:
        return f"{ctx.group(1)}년 {ctx.group(2)}월 {ctx.group(3)}일"
    m = re.search(r'(\d{4})\s*[년./\-]\s*(\d{1,2})\s*[월./\-]\s*(\d{1,2})', text)
    if m:
        return f"{m.group(1)}년 {m.group(2)}월 {m.group(3)}일"
    return ""


def resolve_in_row(row) -> str:
    text = " ".join(str(c) for c in row if c)
    for kw in RESOLVE_KW:
        if kw in text:
            m = re.search(kw + r'[^\s\n]*', text)
            return m.group(0) if m else kw
    return ""


def parse_tables(tables: list) -> list:
    agendas = []
    for table in tables:
        if not table:
            continue
        col_ag, col_rs, col_ct, col_rk = -1, -1, -1, -1
        hdr_idx = -1

        for ri, row in enumerate(table[:5]):
            if not row:
                continue
            joined = " ".join(str(c) for c in row if c)
            if any(k in joined for k in AGENDA_KW | RESULT_KW):
                hdr_idx = ri
                for ci, cell in enumerate(row):
                    s = str(cell) if cell else ""
                    if any(k in s for k in AGENDA_KW):
                        col_ag = ci
                    if any(k in s for k in RESULT_KW):
                        col_rs = ci
                    if any(k in s for k in CONTENT_KW):
                        col_ct = ci
                    if any(k in s for k in REMARK_KW):
                        col_rk = ci
                break

        if hdr_idx < 0 or col_ag < 0:
            continue

        for row in table[hdr_idx + 1:]:
            if not row:
                continue

            def _get(ci):
                return str(row[ci]).strip() if 0 <= ci < len(row) and row[ci] else ""

            name    = _get(col_ag)
            result  = _get(col_rs)
            content = _get(col_ct)
            remark  = _get(col_rk)

            if not name or name in {"None", "-"}:
                continue
            if not result:
                result = resolve_in_row(row)

            agendas.append({"name": name, "content": content, "result": result, "remark": remark})

    return agendas


def parse_text(text: str) -> list:
    agendas = []
    lines = text.split("\n")
    agenda_re = [
        re.compile(r'(?:제|안건)?\s*(\d+)\s*(?:호|번)?\s*안건\s*[:\.\s]\s*(.{3,})'),
        re.compile(r'안건\s*(\d+)\s*[:\.\s]\s*(.{3,})'),
        re.compile(r'^(\d+)\.\s+(.{5,60})$'),
    ]
    for i, line in enumerate(lines):
        line = line.strip()
        for pat in agenda_re:
            m = pat.match(line)
            if m:
                name   = m.group(2).strip()
                result = ""
                content_lines = []

                for j in range(i + 1, min(i + 12, len(lines))):
                    sl = lines[j].strip()
                    if not sl:
                        continue
                    # 의결결과 키워드 발견 시 캡처 후 중단
                    hit = False
                    for kw in RESOLVE_KW:
                        if kw in sl:
                            rm = re.search(kw + r'[^\n]*', sl)
                            result = rm.group(0).strip() if rm else kw
                            hit = True
                            break
                    if hit:
                        break
                    # 다음 안건 시작 시 중단
                    if any(p.match(sl) for p in agenda_re):
                        break
                    content_lines.append(sl)

                content = ' '.join(content_lines[:3])  # 최대 3줄
                agendas.append({"name": name, "content": content, "result": result, "remark": ""})
                break
    return agendas


def extract_from_pdf(pdf_path: Path) -> list:
    results = []
    with pdfplumber.open(pdf_path) as pdf:
        full_text = ""
        all_tables = []
        for page in pdf.pages:
            full_text += (page.extract_text() or "") + "\n"
            tables = page.extract_tables()
            if tables:
                all_tables.extend(tables)

    mtg_name = meeting_name(full_text)
    date_str  = meeting_date(full_text)
    agendas   = parse_tables(all_tables)
    if not agendas:
        agendas = parse_text(full_text)

    if agendas:
        for ag in agendas:
            results.append({
                "파일명":        pdf_path.name,
                "회의명":        mtg_name,
                "회의일":        date_str,
                "안건명":        ag.get("name", ""),
                "안건 주요사항": ag.get("content", ""),
                "의결결과":      ag.get("result", ""),
                "비고":          ag.get("remark", ""),
            })
    else:
        results.append({
            "파일명":        pdf_path.name,
            "회의명":        mtg_name,
            "회의일":        date_str,
            "안건명":        "(안건 추출 불가)",
            "안건 주요사항": "",
            "의결결과":      "",
            "비고":          "",
        })
    return results


# ── 저장 로직 ──────────────────────────────────────────────────

def cell_shading(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def save_excel(data: list, output_dir: Path) -> Path:
    df   = pd.DataFrame(data)
    path = output_dir / "회의록_분석결과.xlsx"
    COLS   = ["파일명", "회의명", "회의일", "안건명", "안건 주요사항", "의결결과", "비고"]
    WIDTHS = [26, 22, 14, 32, 42, 14, 12]

    with pd.ExcelWriter(str(path), engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="회의록", index=False)
        wb = writer.book
        ws = writer.sheets["회의록"]

        hdr_fmt  = wb.add_format({"bold": True, "font_size": 11, "bg_color": "#2E75B6", "font_color": "#FFFFFF", "align": "center", "valign": "vcenter", "border": 1})
        row_odd  = wb.add_format({"font_size": 10, "valign": "vcenter", "border": 1, "text_wrap": True})
        row_even = wb.add_format({"font_size": 10, "valign": "vcenter", "border": 1, "text_wrap": True, "bg_color": "#EBF3FB"})
        pass_fmt = wb.add_format({"font_size": 10, "align": "center", "valign": "vcenter", "border": 1, "bold": True, "font_color": "#196B24", "bg_color": "#E2EFDA"})
        fail_fmt = wb.add_format({"font_size": 10, "align": "center", "valign": "vcenter", "border": 1, "bold": True, "font_color": "#9C0006", "bg_color": "#FFC7CE"})

        for ci, (col, w) in enumerate(zip(COLS, WIDTHS)):
            ws.write(0, ci, col, hdr_fmt)
            ws.set_column(ci, ci, w)

        res_ci = COLS.index("의결결과")
        for ri in range(len(df)):
            base = row_even if ri % 2 == 0 else row_odd
            for ci, col in enumerate(COLS):
                val = df.iloc[ri][col] if col in df.columns else ""
                if ci == res_ci:
                    vs = str(val)
                    if any(k in vs for k in ["가결", "승인"]):
                        ws.write(ri + 1, ci, val, pass_fmt)
                    elif any(k in vs for k in ["부결", "보류"]):
                        ws.write(ri + 1, ci, val, fail_fmt)
                    else:
                        ws.write(ri + 1, ci, val, base)
                else:
                    ws.write(ri + 1, ci, val, base)
            ws.set_row(ri + 1, 20)

        ws.set_row(0, 25)
        ws.freeze_panes(1, 0)

    return path


def save_word(data: list, output_dir: Path) -> Path:
    doc  = Document()
    path = output_dir / "회의록_분석결과.docx"

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("입주자대표회의 회의록 분석 보고서")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    df   = pd.DataFrame(data)
    smry = doc.add_paragraph()
    smr  = smry.add_run(f"■ 분석 요약: 총 {len(df)}건의 안건  /  {df['파일명'].nunique()}개 파일")
    smr.bold = True
    smr.font.size = Pt(11)
    doc.add_paragraph()

    HEADERS  = ["파일명", "회의명", "회의일", "안건명", "안건 주요사항", "의결결과", "비고"]
    COL_W_CM = [3.5, 3.0, 2.8, 3.8, 5.0, 2.4, 1.5]

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"

    for ci, h in enumerate(HEADERS):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell_shading(cell, "2E75B6")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    res_ci = HEADERS.index("의결결과")
    for ri, row_data in enumerate(data):
        row_cells = table.add_row().cells
        values    = [row_data.get(h, "") for h in HEADERS]
        bg = "EBF3FB" if ri % 2 == 0 else None

        for ci, val in enumerate(values):
            cell = row_cells[ci]
            cell.text = str(val) if val else ""
            if bg:
                cell_shading(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if ci == res_ci:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.bold = True
                        vs = str(val)
                        if any(k in vs for k in ["가결", "승인"]):
                            run.font.color.rgb = RGBColor(0x19, 0x6B, 0x24)
                        elif any(k in vs for k in ["부결", "보류"]):
                            run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

    for row in table.rows:
        for ci, w in enumerate(COL_W_CM):
            row.cells[ci].width = Cm(w)

    doc.save(str(path))
    return path


# ── 백그라운드 작업 스레드 ──────────────────────────────────────

def run_analysis(job_id: str, folder_path: str):
    job        = jobs[job_id]
    q          = job['queue']
    output_dir = Path(folder_path)   # 결과 파일을 원본 폴더에 저장

    def log(msg):
        q.put({'type': 'log', 'msg': msg})

    def progress(val):
        q.put({'type': 'progress', 'value': val})

    try:
        pdf_files = sorted(Path(folder_path).glob("*.pdf"))
        if not pdf_files:
            q.put({'type': 'error', 'msg': '폴더에 PDF 파일이 없습니다.'})
            return

        all_data = []
        total    = len(pdf_files)
        log(f"총 {total}개의 PDF 파일을 발견했습니다.\n")

        for i, pdf_path in enumerate(pdf_files):
            progress(int(i / total * 80))
            log(f"[{i+1}/{total}] {pdf_path.name} 처리 중...")
            try:
                rows = extract_from_pdf(pdf_path)
                all_data.extend(rows)
                log(f"  → {len(rows)}건 안건 추출 완료")
            except Exception as e:
                log(f"  → 오류: {e}")

        log(f"\n총 {len(all_data)}건 데이터 추출 완료.")

        log("Excel 파일 저장 중...")
        excel_path = save_excel(all_data, output_dir)
        log(f"[Excel] 저장 완료: {excel_path}")

        log("Word 파일 저장 중...")
        word_path = save_word(all_data, output_dir)
        log(f"[Word] 저장 완료: {word_path}")

        progress(100)
        log("\n모든 작업이 완료되었습니다!")

        job['status']     = 'done'
        job['output_dir'] = output_dir
        q.put({'type': 'done', 'data': all_data})

    except Exception as e:
        log(f"[오류] {e}")
        job['status'] = 'error'
        q.put({'type': 'error', 'msg': str(e)})
    finally:
        q.put(None)


# ── 라우트 ────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/validate', methods=['POST'])
def validate():
    """입력된 폴더 경로가 유효한지 확인하고 PDF 개수를 반환"""
    data = request.get_json()
    if not data or not data.get('folder'):
        return jsonify({'valid': False, 'error': '경로를 입력해주세요.'})

    folder = Path(data['folder'].strip())

    if not folder.exists():
        return jsonify({'valid': False, 'error': '존재하지 않는 폴더입니다.'})
    if not folder.is_dir():
        return jsonify({'valid': False, 'error': '폴더 경로가 아닙니다.'})

    pdf_count = len(list(folder.glob('*.pdf')))
    if pdf_count == 0:
        return jsonify({'valid': False, 'error': 'PDF 파일이 없습니다.'})

    return jsonify({'valid': True, 'pdf_count': pdf_count})


@app.route('/analyze', methods=['POST'])
def analyze():
    data = request.get_json()
    if not data or not data.get('folder'):
        return jsonify({'error': '폴더 경로를 입력해주세요.'}), 400

    folder_path = data['folder'].strip()
    folder      = Path(folder_path)

    if not folder.exists():
        return jsonify({'error': f'폴더가 존재하지 않습니다: {folder_path}'}), 400
    if not folder.is_dir():
        return jsonify({'error': f'폴더 경로가 아닙니다: {folder_path}'}), 400

    pdf_count = len(list(folder.glob("*.pdf")))
    if pdf_count == 0:
        return jsonify({'error': '폴더에 PDF 파일이 없습니다.'}), 400

    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        'queue':      queue.Queue(),
        'status':     'running',
        'output_dir': folder,
    }

    t = threading.Thread(target=run_analysis, args=(job_id, str(folder)), daemon=True)
    t.start()

    return jsonify({'job_id': job_id, 'pdf_count': pdf_count})


@app.route('/stream/<job_id>')
def stream(job_id):
    if job_id not in jobs:
        return jsonify({'error': '잘못된 job_id'}), 404

    def generate():
        q = jobs[job_id]['queue']
        while True:
            item = q.get()
            if item is None:
                break
            yield f"data: {json.dumps(item, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


@app.route('/download/<job_id>/<filetype>')
def download(job_id, filetype):
    if job_id not in jobs:
        return jsonify({'error': '잘못된 job_id'}), 404

    output_dir = jobs[job_id]['output_dir']

    if filetype == 'excel':
        path = output_dir / '회의록_분석결과.xlsx'
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    elif filetype == 'word':
        path = output_dir / '회의록_분석결과.docx'
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        return jsonify({'error': '잘못된 파일 형식'}), 400

    if not path.exists():
        return jsonify({'error': '파일이 없습니다. 먼저 분석을 실행해주세요.'}), 404

    return send_file(str(path), as_attachment=True, mimetype=mimetype)


if __name__ == '__main__':
    app.run(debug=True, port=5000, threaded=True)
